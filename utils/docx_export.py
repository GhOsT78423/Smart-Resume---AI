from docx import Document
import os

class DOCXExporter:
    """
    Exports resumes as DOCX.
    """

    @staticmethod
    def export(content: str, filename: str):

        os.makedirs("exports/docx", exist_ok=True)

        document = Document()

        document.add_heading("Resume", level=1)

        for line in content.split("\n"):
            document.add_paragraph(line)

        output_path = f"exports/docx/{filename}.docx"

        document.save(output_path)

        return output_path